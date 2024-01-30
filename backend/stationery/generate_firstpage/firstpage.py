import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_run_font(run, font_name='Times New Roman', font_size=Pt(16), bold=False):
    font = run.font
    font.name = font_name
    font.size = font_size
    font.bold = bold

def create_word_file(subject_name, subject_code, faculty_name, student_name,
                     faculty_designation, roll_number, semester, group, image_path):

    # Create a new Word document
    doc = Document()

    # Add a centered heading for Subject
    subject_heading = doc.add_paragraph()
    subject_run = subject_heading.add_run(f'{subject_name.upper()}')
    set_run_font(subject_run, font_size=Pt(20), bold=True)
    subject_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a centered heading for Subject Code
    subject_code_heading = doc.add_paragraph()
    subject_code_run = subject_code_heading.add_run(f'({subject_code.upper()})')
    set_run_font(subject_code_run, font_size=Pt(20), bold=True)
    subject_code_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add 3 blank lines
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Add subheading for Faculty Name and Student Name in one line
    faculty_student_subheading = doc.add_paragraph()
    
    faculty_run = faculty_student_subheading.add_run('Faculty Name:')
    set_run_font(faculty_run, font_size=Pt(18), bold=True)
    faculty_student_subheading.add_run(' ' * (83))

    student_run = faculty_student_subheading.add_run('Student Name:')
    set_run_font(student_run, font_size=Pt(18), bold=True)
    faculty_student_subheading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Calculate the number of spaces needed based on the length of the faculty name
    if ( (len(student_name) + len(faculty_name)) < 15 ):
        spaces_needed = 75 - len(faculty_name)
        
    elif ( (len(student_name) + len(faculty_name)) >= 15 and (len(student_name) + len(faculty_name)) <= 22):
        spaces_needed = 68 - len(faculty_name)
        
    elif ( (len(student_name) + len(faculty_name)) > 22 and (len(student_name) + len(faculty_name)) <= 30):
        spaces_needed = 60 - len(faculty_name)
        
    elif ( (len(student_name) + len(faculty_name)) > 30 and (len(student_name) + len(faculty_name)) < 37):
        if ( len(faculty_name)/(len(faculty_name)+len(student_name)) > 0.6):
            spaces_needed = 56 - len(faculty_name)
        else:
            spaces_needed = 52 - len(faculty_name)
        
    else :
        spaces_needed = 45 - len(faculty_name)

    # Create a string with the required spaces between faculty and student names
    formatted_line = f"{faculty_name}{' ' * spaces_needed}{student_name}"

    # Add the line with faculty_name and student_name
    faculty_student_names_line = doc.add_paragraph()
    faculty_name_run = faculty_student_names_line.add_run(formatted_line)
    set_run_font(faculty_name_run)
    faculty_student_names_line.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE

    # Add the faculty designation and Roll No
    faculty_designation_line = doc.add_paragraph()
    faculty_designation_run = faculty_designation_line.add_run(f"({faculty_designation}){' '  * 35}")
    set_run_font(faculty_designation_run)
    roll_number_run = faculty_designation_line.add_run(f'Roll No: {roll_number}')
    set_run_font(roll_number_run, font_size=Pt(14))
    faculty_designation_line.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
    
    # Add right-aligned text for Semester
    semester_line = doc.add_paragraph()
    semester_run = semester_line.add_run(f'Semester: {semester}')
    set_run_font(semester_run, font_size=Pt(14))
    semester_line.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add right-aligned text for Group
    group_line = doc.add_paragraph()
    group_run = group_line.add_run(f'Group: {group}')
    set_run_font(group_run, font_size=Pt(14))
    group_line.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Leave 2 lines blank
    doc.add_paragraph()
    doc.add_paragraph()

    # Insert the image and center-align it
    image_absolute_path = str(Path(__file__).resolve().parent / image_path)
    image_paragraph = doc.add_paragraph()
    image_run = image_paragraph.add_run()
    image_run.add_picture(image_absolute_path, width=Inches(1.8), height=Inches(1.8))
    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Leave 1 line blank
    doc.add_paragraph()

    # Add centered text for the institute's address
    institute_paragraph_1 = doc.add_paragraph()
    institute_run_1 = institute_paragraph_1.add_run('Maharaja Agrasen Institute of Technology, PSP Area,')
    set_run_font(institute_run_1, font_size=Pt(14))
    institute_paragraph_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    institute_paragraph_2 = doc.add_paragraph()
    institute_run_2 = institute_paragraph_2.add_run('Sector – 22, Rohini, New Delhi -110085')
    set_run_font(institute_run_2, font_size=Pt(14))
    institute_paragraph_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    output_filename = str(Path(__file__).resolve().parent / 'first_page.docx')

    # Check if the file already exists, and delete it
    if os.path.exists(output_filename):
        os.remove(output_filename)

    # Save the document
    doc.save(output_filename)
    
    return output_filename

'''
# Example usage:
create_word_file(subject_name='Programming in Java Lab', subject_code='CIC-258', faculty_name='Ms. Kajol',
                 student_name='Chanmeet Singh Sahni', faculty_designation='Assistant Professor', roll_number='01296402722',
                 semester='4th', group='3C11', image_path='maitlogomain.png')
'''